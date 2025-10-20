#!/usr/bin/env python3
"""
Worker module for file conversion
Handles the actual conversion logic
"""

import os
import re
from dataclasses import dataclass
from typing import List, Any

import pytesseract
from docx import Document
import pdfplumber
from openpyxl import load_workbook
import html2text
import chardet
from PIL import Image


@dataclass
class ConversionOptions:
    """Configuration options for conversion"""
    ocr_enabled: bool = True
    compress_images: bool = True
    extract_tables: bool = True
    add_metadata: bool = False
    output_folder: str = "."


class FileConverter:
    """Handles file to markdown conversion"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf', '.docx', '.doc', '.xlsx', '.xls',
            '.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff',
            '.txt', '.html', '.htm', '.csv'
        }
    
    def check_dependencies(self) -> List[str]:
        """Check which dependencies are missing"""
        missing: List[str] = []
        
        # Check Tesseract executable
        try:
            pytesseract.get_tesseract_version()
        except Exception:
            missing.append("Tesseract OCR (executable not found in PATH)")
        
        return missing
    
    def convert_file(self, input_path: str, options: ConversionOptions) -> str:
        """
        Convert a file to markdown
        
        Args:
            input_path: Path to input file
            options: Conversion options
            
        Returns:
            Path to output markdown file
            
        Raises:
            ValueError: If file type is unsupported
            FileNotFoundError: If input file doesn't exist
        """
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        ext = os.path.splitext(input_path)[1].lower()
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported file type: {ext}")
        
        # Determine conversion method
        if ext == '.pdf':
            markdown = self._convert_pdf(input_path, options)
        elif ext in {'.docx', '.doc'}:
            markdown = self._convert_docx(input_path, options)
        elif ext in {'.xlsx', '.xls'}:
            markdown = self._convert_xlsx(input_path, options)
        elif ext in {'.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff'}:
            markdown = self._convert_image(input_path, options)
        elif ext in {'.html', '.htm'}:
            markdown = self._convert_html(input_path, options)
        elif ext == '.txt':
            markdown = self._convert_txt(input_path, options)
        elif ext == '.csv':
            markdown = self._convert_csv(input_path, options)
        else:
            raise ValueError(f"No converter for: {ext}")
        
        # Add metadata if requested
        if options.add_metadata:
            metadata = self._generate_metadata(input_path)
            markdown = f"{metadata}\n\n{markdown}"
        
        # Save to output folder
        output_path = self._get_output_path(input_path, options.output_folder)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown)
        
        return output_path
    
    def _get_output_path(self, input_path: str, output_folder: str) -> str:
        """Generate output file path"""
        basename = os.path.splitext(os.path.basename(input_path))[0]
        return os.path.join(output_folder, f"{basename}.md")
    
    def _generate_metadata(self, input_path: str) -> str:
        """Generate YAML frontmatter"""
        filename = os.path.basename(input_path)
        file_size = os.path.getsize(input_path)
        
        return f"""---
source_file: {filename}
size_bytes: {file_size}
converted_with: openElara File Converter
---"""
    
    def _convert_pdf(self, input_path: str, options: ConversionOptions) -> str:
        """Convert PDF to markdown"""
        
        content: List[str] = []
        
        with pdfplumber.open(input_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Try text extraction first
                text = page.extract_text(layout=True)
                
                if text and text.strip():
                    content.append(text)
                elif options.ocr_enabled:
                    # Fallback to OCR if no text found
                    try:
                        img = page.to_image(resolution=150 if options.compress_images else 300)
                        ocr_text = pytesseract.image_to_string(img.original)
                        content.append(ocr_text)
                    except Exception as e:
                        content.append(f"[OCR failed on page {page_num}: {str(e)}]")
                else:
                    content.append(f"[Page {page_num}: No text extracted]")
        
        return self._clean_text("\n\n".join(content))
    
    def _convert_docx(self, input_path: str, options: ConversionOptions) -> str:
        """Convert DOCX to markdown"""
        
        doc = Document(input_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract tables if requested
        if options.extract_tables and doc.tables:
            for table in doc.tables:
                table_md = self._table_to_markdown(table)
                paragraphs.append(table_md)
        
        return self._clean_text("\n\n".join(paragraphs))
    
    def _convert_xlsx(self, input_path: str, options: ConversionOptions) -> str:
        """Convert XLSX to markdown"""
        
        workbook = load_workbook(filename=input_path)
        sheets_md: List[str] = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            if sheet.max_row > 0:
                sheets_md.append(f"## {sheet_name}\n")
                
                # Get header (first row)
                header: List[str] = [str(cell.value) if cell.value is not None else "" for cell in sheet[1]]
                
                # Get data rows
                rows: List[List[str]] = []
                for row in sheet.iter_rows(min_row=2, values_only=True):
                    rows.append([str(cell) if cell is not None else "" for cell in row])
                
                table_md = self._table_to_markdown_list(header, rows)
                sheets_md.append(table_md)
        
        return "\n\n".join(sheets_md)
    
    def _convert_image(self, input_path: str, options: ConversionOptions) -> str:
        """Convert image to markdown (OCR)"""
        
        # Load image
        img = Image.open(input_path)
        
        # Compress if requested
        if options.compress_images:
            # Resize if too large
            max_size = (2000, 2000)
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
        
        # Run OCR
        text = pytesseract.image_to_string(img)
        
        return self._clean_text(text)
    
    def _convert_html(self, input_path: str, options: ConversionOptions) -> str:
        """Convert HTML to markdown"""
        
        with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
            html_content = f.read()
        
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = False
        
        return h.handle(html_content)
    
    def _convert_txt(self, input_path: str, options: ConversionOptions) -> str:
        """Convert TXT to markdown"""
        # Detect encoding
        encoding = 'utf-8'
        with open(input_path, 'rb') as f:
            raw_data = f.read()
            detected = chardet.detect(raw_data)
            if detected['confidence'] > 0.5:
                encoding = detected['encoding']
        
        with open(input_path, 'r', encoding=encoding, errors='replace') as f:
            return f.read()
    
    def _convert_csv(self, input_path: str, options: ConversionOptions) -> str:
        """Convert CSV to markdown table"""
        import csv
        
        with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
            reader = csv.reader(f)
            rows = list(reader)
        
        if not rows:
            return ""
        
        header = rows[0]
        data_rows = rows[1:]
        
        return self._table_to_markdown_list(header, data_rows)
    
    def _table_to_markdown(self, table: Any) -> str:
        """Convert DOCX table to markdown"""
        rows: List[List[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        
        if not rows:
            return ""
        
        header = rows[0]
        data_rows = rows[1:]
        
        return self._table_to_markdown_list(header, data_rows)
    
    def _table_to_markdown_list(self, header: List[str], rows: List[List[str]]) -> str:
        """Convert header and rows to markdown table"""
        if not header:
            return ""
        
        # Build markdown table
        md = "| " + " | ".join(str(h) for h in header) + " |\n"
        md += "|" + "|".join(["---"] * len(header)) + "|\n"
        
        for row in rows:
            # Pad row if shorter than header
            while len(row) < len(header):
                row.append("")
            md += "| " + " | ".join(str(cell) for cell in row[:len(header)]) + " |\n"
        
        return md
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Fix missing spaces
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', text)
        text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', text)
        text = re.sub(r'([.,!?])([a-zA-Z])', r'\1 \2', text)
        
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
