import sys
import json
import os
import re
import traceback
from typing import Dict, Any, List

try:
    from docx import Document
    from docx.shared import Pt
    docx_available = True
except ImportError:
    docx_available = False


def parse_markdown_to_docx(markdown_text: str, docx_path: str) -> Dict[str, Any]:
    try:
        if not docx_available:
            return {
                'success': False,
                'error': 'python-docx not available. Install: pip install python-docx'
            }
        
        doc = Document()  # type: ignore
        lines: List[str] = markdown_text.split('\n')
        
        in_code_block: bool = False
        code_lines: List[str] = []
        in_list: bool = False
        
        for line in lines:
            if line.strip().startswith('```'):
                if in_code_block:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = 'Intense Quote'
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            
            if in_code_block:
                code_lines.append(line)
                continue
            
            if line.startswith('#'):
                in_list = False
                heading_level = len(line) - len(line.lstrip('#'))
                heading_text = line.lstrip('#').strip()
                
                if heading_level == 1:
                    p = doc.add_heading(heading_text, level=1)
                elif heading_level == 2:
                    p = doc.add_heading(heading_text, level=2)
                else:
                    p = doc.add_heading(heading_text, level=3)
                continue
            
            if line.strip().startswith(('-', '*', '•')):
                text = line.strip().lstrip('-*•').strip()
                p = doc.add_paragraph(text, style='List Bullet')
                in_list = True
                continue
            
            numbered_match = re.match(r'^(\d+)[\.\)]\s+(.*)', line.strip())
            if numbered_match:
                text = numbered_match.group(2)
                p = doc.add_paragraph(text, style='List Number')
                in_list = True
                continue
            
            if line.strip().startswith('>'):
                in_list = False
                text = line.strip().lstrip('>').strip()
                p = doc.add_paragraph(text, style='Quote')
                continue
            
            if line.strip() in ('---', '***', '___'):
                in_list = False
                doc.add_paragraph('_' * 50)
                continue
            
            if not line.strip():
                if not in_list:
                    doc.add_paragraph()
                continue
            
            in_list = False
            p = doc.add_paragraph()
            
            text = line
            
            parts = re.split(r'`([^`]+)`', text)
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)  # type: ignore
                else:
                    sub_parts = re.split(r'\*\*([^\*]+)\*\*|\*([^\*]+)\*', part)
                    for j, sub_part in enumerate(sub_parts):
                        if sub_part:
                            run = p.add_run(sub_part)
                            if j % 3 == 1:
                                run.bold = True
                            elif j % 3 == 2:
                                run.italic = True
        
        os.makedirs(os.path.dirname(docx_path), exist_ok=True)
        
        doc.save(docx_path)
        
        return {
            'success': True,
            'output_path': docx_path,
            'message': 'Successfully converted Markdown to DOCX'
        }
    
    except Exception as e:
        return {
            'success': False,
            'error': f'Markdown to DOCX conversion failed: {str(e)}'
        }


def markdown_file_to_docx(md_path: str, docx_path: str) -> Dict[str, Any]:
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        
        return parse_markdown_to_docx(markdown_text, docx_path)
    
    except FileNotFoundError:
        return {'success': False, 'error': f'Markdown file not found: {md_path}'}
    except Exception as e:
        return {'success': False, 'error': f'Failed to read Markdown file: {str(e)}'}


def main():
    try:
        input_data = json.loads(sys.stdin.read())
        
        input_path = input_data.get('input')
        output_path = input_data.get('output')
        
        if not input_path or not output_path:
            print(json.dumps({
                'success': False,
                'error': 'Missing input or output path'
            }))
            return
        
        result = markdown_file_to_docx(input_path, output_path)
        print(json.dumps(result), flush=True)
    
    except json.JSONDecodeError as e:
        print(json.dumps({
            'success': False,
            'error': f'Invalid JSON input: {str(e)}'
        }), file=sys.stderr, flush=True)
    
    except Exception as e:
        print(json.dumps({
            'success': False,
            'error': f'Unexpected error: {str(e)}',
            'traceback': traceback.format_exc()
        }), file=sys.stderr, flush=True)


if __name__ == '__main__':
    main()
